import docx, sys, os

def extract_to_text(docx_path, out_path):
    doc = docx.Document(docx_path)
    lines = []
    lines.append('=' * 60)
    lines.append('文件名: ' + os.path.basename(docx_path))
    lines.append('=' * 60)
    
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    
    for t_idx, table in enumerate(doc.tables):
        lines.append('')
        lines.append('--- 表格 ' + str(t_idx + 1) + ' (' + str(len(table.rows)) + '行 x ' + str(len(table.columns)) + '列) ---')
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            lines.append('  R' + str(r_idx) + ': ' + ' | '.join(cells))
    
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return len(lines)

n1 = extract_to_text(r'C:\Users\Administrator\Desktop\2026年度团体标准计划项目汇总表.docx', r'C:\Users\Administrator\Desktop\doc1_extracted.txt')
n2 = extract_to_text(r'C:\Users\Administrator\Desktop\0323CAEE-2026年度团体标准计划项目汇总表(4).docx', r'C:\Users\Administrator\Desktop\doc2_extracted.txt')
print('文件1: ' + str(n1) + '行')
print('文件2: ' + str(n2) + '行')
