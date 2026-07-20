# -*- coding: utf-8 -*-
import docx
import sys

def extract(path, outpath):
    doc = docx.Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for t_idx, table in enumerate(doc.tables):
        lines.append("")
        lines.append("=== 表格 {} ({}行 x {}列) ===".format(t_idx+1, len(table.rows), len(table.columns)))
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            lines.append("R{}: {}".format(r_idx, " || ".join(cells)))
    with open(outpath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return len(lines)

n1 = extract(
    r"C:\Users\Administrator\Desktop\2026年度团体标准计划项目汇总表.docx",
    r"C:\Users\Administrator\Desktop\doc1_out.txt"
)
n2 = extract(
    r"C:\Users\Administrator\Desktop\0323CAEE-2026年度团体标准计划项目汇总表(4).docx",
    r"C:\Users\Administrator\Desktop\doc2_out.txt"
)
print("文件1: {}行".format(n1))
print("文件2: {}行".format(n2))
